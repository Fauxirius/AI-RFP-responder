import docx
from docx.shared import Pt
import io

def export_to_docx(offer_sections):
    """
    Compiles the offer sections into a Word document.
    """
    doc = docx.Document()
    
    # Title
    doc.add_heading('Generated Offer', 0)
    
    for section_name, content in offer_sections.items():
        doc.add_heading(section_name, level=1)
        # Basic markdown cleanup for the content
        # Note: A full markdown-to-docx converter would be better for complex formatting,
        # but for this MVP we'll just dump the text and handle basic paragraphs.
        
        # Split by newlines to create paragraphs
        for paragraph in content.split('\n'):
            if paragraph.strip():
                if paragraph.startswith('## '):
                    doc.add_heading(paragraph.replace('## ', ''), level=2)
                elif paragraph.startswith('### '):
                    doc.add_heading(paragraph.replace('### ', ''), level=3)
                elif paragraph.startswith('- '):
                    doc.add_paragraph(paragraph.replace('- ', ''), style='List Bullet')
                elif paragraph.startswith('* '):
                    doc.add_paragraph(paragraph.replace('* ', ''), style='List Bullet')
                else:
                    doc.add_paragraph(paragraph)
    
    # Save to a byte stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream
